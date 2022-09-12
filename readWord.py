from docx import Document
import movingFile



filename = 'allFiles/АО «Астраханьгазсервис» 10.09.2017.doc'
doc = Document(filename)

all_tables = doc.tables
# print('Всего таблиц в документе:', len(all_tables))
count=0
# создаем пустой словарь под данные таблиц
data_tables = {i:None for i in range(len(all_tables))}
# проходимся по таблицам
for i, table in enumerate(all_tables):
    # print('\nДанные таблицы №', i)
    # создаем список строк для таблицы `i` (пока пустые)
    data_tables[i] = [[] for _ in range(len(table.rows))]
    # проходимся по строкам таблицы `i`
    for j, row in enumerate(table.rows):
        # проходимся по ячейкам таблицы `i` и строки `j`
        for cell in row.cells:
            # добавляем значение ячейки в соответствующий
            # список, созданного словаря под данные таблиц
            data_tables[i][j].append(cell.text)
        count+=1
    # смотрим извлеченные данные
    # (по строкам) для таблицы `i`
    # print(data_tables[i])
    flag=0
    for key in data_tables[0]:
        for word in key:
            if 'происшествия' in word:
                movingFile.move(filename)
                flag=1
                break
        if flag:
            break
    print('\n')

# print('Данные всех таблиц документа:')
# print(data_tables)













# for table in document.tables:
#     for row in (table.rows[1],table.rows[-1]):
#         for cell in row.cells:
#             print (cell.text)